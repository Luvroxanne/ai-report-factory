from __future__ import annotations

import math
from pathlib import Path
from typing import Any


class ScriptService:
    def build_scripts(self, plan: dict[str, Any]) -> list[dict[str, Any]]:
        scripts: list[dict[str, Any]] = []
        for index, slide in enumerate(plan.get("slides", []), start=1):
            title = str(slide.get("title") or f"第 {index} 页")
            bullets = [str(item) for item in (slide.get("bullets") or []) if str(item).strip()]
            note = str(slide.get("speaker_note") or "").strip()
            if not note:
                point_text = "；".join(bullets[:4]) if bullets else "本页为过渡说明"
                note = f"接下来介绍{title}。本页重点包括：{point_text}。讲解时可以先给出结论，再说明关键原因和落地动作。"
            estimated_seconds = int(slide.get("estimated_seconds") or self._estimate_seconds(note))
            scripts.append(
                {
                    "page": index,
                    "title": title,
                    "chapter": str(slide.get("chapter") or ""),
                    "layout": str(slide.get("layout") or "content"),
                    "text": note,
                    "estimated_seconds": estimated_seconds,
                    "remarks": self._remarks(slide),
                }
            )
        return scripts

    def save(self, scripts: list[dict[str, Any]], output_path: Path, report_title: str = "AI报告解说稿") -> Path:
        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")
        return self._save_docx(scripts, output_path, report_title)

    def _save_docx(self, scripts: list[dict[str, Any]], output_path: Path, report_title: str) -> Path:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)

        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(report_title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(20, 64, 128)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("按 PPT 页码组织的口播稿，包含预计讲解时长与备注。")
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.color.rgb = RGBColor(102, 120, 145)

        total_seconds = sum(int(item.get("estimated_seconds") or 0) for item in scripts)
        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"共 {len(scripts)} 页 · 预计总时长 {self._format_duration(total_seconds)}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(120, 136, 156)

        document.add_paragraph("")

        for item in scripts:
            heading = document.add_paragraph()
            heading.style = styles["Heading 1"]
            heading_run = heading.add_run(f"第 {item['page']} 页：{item['title']}")
            heading_run.font.name = "Microsoft YaHei"
            heading_run.font.size = Pt(15)
            heading_run.font.color.rgb = RGBColor(28, 88, 168)

            info = document.add_paragraph()
            info_run = info.add_run(
                f"章节：{item.get('chapter') or '未分组'} · 版式：{item.get('layout') or 'content'} · 预计时长：{self._format_duration(int(item.get('estimated_seconds') or 0))}"
            )
            info_run.font.size = Pt(9)
            info_run.font.color.rgb = RGBColor(110, 126, 148)

            para = document.add_paragraph()
            para.paragraph_format.line_spacing = 1.45
            para.paragraph_format.space_after = Pt(10)
            body = para.add_run(str(item["text"]))
            body.font.name = "Microsoft YaHei"
            body.font.size = Pt(11)
            body.font.color.rgb = RGBColor(40, 48, 62)

            remarks = document.add_paragraph()
            remarks_run = remarks.add_run(f"备注：{item.get('remarks') or '保持自然语速，避免照读屏幕文字。'}")
            remarks_run.font.size = Pt(9)
            remarks_run.font.italic = True
            remarks_run.font.color.rgb = RGBColor(126, 138, 155)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    def _estimate_seconds(self, text: str) -> int:
        # 估算中文口播约 4.5-5 字/秒，并留出翻页停顿。
        seconds = math.ceil(len(text) / 4.6) + 4
        return max(12, min(95, seconds))

    def _format_duration(self, seconds: int) -> str:
        minutes = seconds // 60
        remain = seconds % 60
        if minutes:
            return f"{minutes}分{remain:02d}秒"
        return f"{remain}秒"

    def _remarks(self, slide: dict[str, Any]) -> str:
        layout = str(slide.get("layout") or "content")
        if layout == "cover":
            return "开场页建议先建立听众预期。"
        if layout == "agenda":
            return "目录页可快速说明整体结构。"
        if layout == "section":
            return "章节页建议停顿 1 秒再展开。"
        if layout == "summary":
            return "总结页强调结论和下一步。"
        return "正文页建议结合案例或数据补充说明。"
